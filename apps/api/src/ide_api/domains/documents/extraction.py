from __future__ import annotations

from io import BytesIO
from pathlib import PurePosixPath
from zipfile import ZipFile

import fitz
from docx import Document as DocxDocument

_MAX_ANALYSIS_CHARACTERS = 200_000
_MAX_ANALYSIS_PDF_PAGES = 10


class DocumentTextExtractionError(Exception):
    pass


def extract_document_text(content: bytes, filename: str) -> str:
    suffix = PurePosixPath(filename).suffix.lower()
    try:
        if suffix == ".docx":
            text = _extract_docx(content)
        elif suffix == ".pdf":
            text = _extract_pdf(content)
        elif suffix == ".tex":
            text = content.decode("utf-8")
        elif suffix == ".zip":
            text = _extract_latex_archive(content)
        else:
            raise DocumentTextExtractionError("Unsupported document type for analysis.")
    except DocumentTextExtractionError:
        raise
    except Exception as error:
        raise DocumentTextExtractionError("Unable to extract document text.") from error

    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if suffix == ".pdf" and not normalized:
        return "[Scanned PDF; page images are attached for visual analysis.]"
    if not normalized:
        raise DocumentTextExtractionError("The document has no extractable text.")
    return normalized[:_MAX_ANALYSIS_CHARACTERS]


def render_scanned_pdf_pages(content: bytes) -> tuple[bytes, ...]:
    document = fitz.open(stream=content, filetype="pdf")
    try:
        if any(page.get_text("text").strip() for page in document):
            return ()
        matrix = fitz.Matrix(1.5, 1.5)
        return tuple(
            page.get_pixmap(matrix=matrix, alpha=False).tobytes("jpeg")
            for page in list(document)[:_MAX_ANALYSIS_PDF_PAGES]
        )
    finally:
        document.close()


def _extract_docx(content: bytes) -> str:
    document = DocxDocument(BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_pdf(content: bytes) -> str:
    document = fitz.open(stream=content, filetype="pdf")
    try:
        parts = []
        for page in document:
            text = page.get_text("text").strip()
            if text:
                parts.append(f"[page {page.number + 1}]\n{text}")
        return "\n".join(parts)
    finally:
        document.close()


def _extract_latex_archive(content: bytes) -> str:
    with ZipFile(BytesIO(content)) as archive:
        parts: list[str] = []
        for name in sorted(archive.namelist()):
            if name.lower().endswith((".tex", ".bib")):
                parts.append(f"[file {name}]\n{archive.read(name).decode('utf-8')}")
        return "\n".join(parts)
