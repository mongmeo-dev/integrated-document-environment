from __future__ import annotations

import asyncio
from collections.abc import Iterable
from io import BytesIO
from typing import Protocol

import fitz
from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table

from ide_api.domains.documents.models import DocumentVersion
from ide_api.domains.formatting.models import ExternalEditResult
from ide_api.domains.formatting.schemas import (
    DetectedFormatDifference,
    FormatDifferenceCategory,
    OriginalFormat,
)
from ide_api.infrastructure.object_storage import ObjectStorage


class ObjectDownloader(Protocol):
    def download(self, object_key: str): ...


class StorageFormatComparisonRunner:
    """Compares immutable source and external-result objects without writing either object."""

    def __init__(self, storage: ObjectDownloader | None = None) -> None:
        self._storage = storage or ObjectStorage()

    async def compare(
        self, *, original: DocumentVersion, result: ExternalEditResult
    ) -> list[DetectedFormatDifference]:
        expected = self._format_for_input_kind(original.input_kind)
        if expected is None or result.original_format != expected.value:
            return [
                self._difference("document", "format", str(expected), str(result.original_format))
            ]

        try:
            original_bytes, result_bytes = await asyncio.gather(
                asyncio.to_thread(self._read_object, original.object_key),
                asyncio.to_thread(self._read_object, result.object_key),
            )
        except Exception as error:
            return [self._difference("document", "storage", "readable", self._error_value(error))]

        if expected is OriginalFormat.DOCX:
            return self._compare_docx(original_bytes, result_bytes)
        return self._compare_pdf(original_bytes, result_bytes)

    def _read_object(self, object_key: str) -> bytes:
        stream = self._storage.download(object_key)
        try:
            return stream.read()
        finally:
            close = getattr(stream, "close", None)
            if close is not None:
                close()

    def _compare_docx(
        self, original_bytes: bytes, result_bytes: bytes
    ) -> list[DetectedFormatDifference]:
        try:
            original = Document(BytesIO(original_bytes))
        except Exception as error:
            return [
                self._difference("document", "docx_parse", "parseable", self._error_value(error))
            ]
        try:
            proposed = Document(BytesIO(result_bytes))
        except Exception as error:
            return [
                self._difference("document", "docx_parse", "parseable", self._error_value(error))
            ]

        differences: list[DetectedFormatDifference] = []
        self._compare_values(
            differences, "sections", self._section_values(original), self._section_values(proposed)
        )
        self._compare_values(
            differences,
            "paragraphs",
            self._paragraph_values(original),
            self._paragraph_values(proposed),
        )
        self._compare_values(
            differences, "tables", self._table_values(original), self._table_values(proposed)
        )
        return differences

    def _section_values(self, document: DocxDocument) -> list[tuple[str, tuple[object, ...]]]:
        return [
            (
                f"section[{index}].margins",
                tuple(
                    self._value(getattr(section, name))
                    for name in ("top_margin", "bottom_margin", "left_margin", "right_margin")
                ),
            )
            for index, section in enumerate(document.sections)
        ]

    def _paragraph_values(self, document: DocxDocument) -> list[tuple[str, tuple[object, ...]]]:
        values: list[tuple[str, tuple[object, ...]]] = []
        for location, paragraph in self._paragraphs(document):
            values.append(
                (
                    f"{location}.line_spacing",
                    (self._value(paragraph.paragraph_format.line_spacing),),
                )
            )
            for run_index, run in enumerate(paragraph.runs):
                font = run.font
                values.extend(
                    (
                        (f"{location}.runs[{run_index}].font.name", (self._value(font.name),)),
                        (f"{location}.runs[{run_index}].font.size", (self._value(font.size),)),
                        (
                            f"{location}.runs[{run_index}].font.color",
                            (self._color_value(font.color.rgb),),
                        ),
                    )
                )
        return values

    def _paragraphs(self, document: DocxDocument) -> Iterable[tuple[str, object]]:
        for index, paragraph in enumerate(document.paragraphs):
            yield f"paragraphs[{index}]", paragraph
        for table_index, table in enumerate(document.tables):
            yield from self._table_paragraphs(table, f"tables[{table_index}]")

    def _table_paragraphs(self, table: Table, location: str) -> Iterable[tuple[str, object]]:
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                cell_location = f"{location}.rows[{row_index}].cells[{cell_index}]"
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    yield f"{cell_location}.paragraphs[{paragraph_index}]", paragraph
                for nested_index, nested in enumerate(cell.tables):
                    yield from self._table_paragraphs(
                        nested, f"{cell_location}.tables[{nested_index}]"
                    )

    def _table_values(self, document: DocxDocument) -> list[tuple[str, tuple[object, ...]]]:
        values: list[tuple[str, tuple[object, ...]]] = []
        for table_index, table in enumerate(document.tables):
            values.extend(self._one_table_values(table, f"tables[{table_index}]"))
        return values

    def _one_table_values(
        self, table: Table, location: str
    ) -> list[tuple[str, tuple[object, ...]]]:
        values = [(f"{location}.structure", (len(table.rows), len(table.columns)))]
        for row_index, row in enumerate(table.rows):
            values.append((f"{location}.rows[{row_index}].cells", (len(row.cells),)))
            for cell_index, cell in enumerate(row.cells):
                cell_location = f"{location}.rows[{row_index}].cells[{cell_index}]"
                values.append(
                    (f"{cell_location}.structure", (len(cell.paragraphs), len(cell.tables)))
                )
                for nested_index, nested in enumerate(cell.tables):
                    values.extend(
                        self._one_table_values(nested, f"{cell_location}.tables[{nested_index}]")
                    )
        return values

    def _compare_pdf(
        self, original_bytes: bytes, result_bytes: bytes
    ) -> list[DetectedFormatDifference]:
        try:
            original = fitz.open(stream=original_bytes, filetype="pdf")
        except Exception as error:
            return [
                self._difference("document", "pdf_parse", "parseable", self._error_value(error))
            ]
        try:
            proposed = fitz.open(stream=result_bytes, filetype="pdf")
        except Exception as error:
            original.close()
            return [
                self._difference("document", "pdf_parse", "parseable", self._error_value(error))
            ]
        try:
            return self._compare_pdf_documents(original, proposed)
        finally:
            original.close()
            proposed.close()

    def _compare_pdf_documents(
        self, original: fitz.Document, proposed: fitz.Document
    ) -> list[DetectedFormatDifference]:
        differences: list[DetectedFormatDifference] = []
        self._compare_values(
            differences,
            "pages",
            [("page_count", (len(original),))],
            [("page_count", (len(proposed),))],
        )
        for page_index in range(max(len(original), len(proposed))):
            original_values = (
                self._pdf_page_values(original[page_index]) if page_index < len(original) else []
            )
            proposed_values = (
                self._pdf_page_values(proposed[page_index]) if page_index < len(proposed) else []
            )
            self._compare_values(
                differences, f"pages[{page_index}]", original_values, proposed_values
            )
            if self._has_no_text_spans(original_values) or self._has_no_text_spans(proposed_values):
                differences.append(
                    self._difference(
                        f"pages[{page_index}].text_spans",
                        "scan_pdf",
                        "extractable text required",
                        "missing",
                    )
                )
        return differences

    def _pdf_page_values(self, page: fitz.Page) -> list[tuple[str, tuple[object, ...]]]:
        page_values: list[tuple[str, tuple[object, ...]]] = [
            ("size", (round(page.rect.width, 4), round(page.rect.height, 4))),
            ("drawing_count", (len(page.get_drawings()),)),
            ("image_count", (len(page.get_images(full=True)),)),
        ]
        spans = [
            span
            for block in page.get_text("dict")["blocks"]
            if block.get("type") == 0
            for line in block["lines"]
            for span in line["spans"]
        ]
        if not spans:
            page_values.append(("text_spans", ("missing",)))
        for span_index, span in enumerate(spans):
            page_values.extend(
                (
                    (f"spans[{span_index}].font", (span.get("font"),)),
                    (f"spans[{span_index}].size", (round(span.get("size", 0), 4),)),
                    (f"spans[{span_index}].color", (span.get("color"),)),
                )
            )
        return page_values

    def _compare_values(
        self,
        differences: list[DetectedFormatDifference],
        prefix: str,
        original: list[tuple[str, tuple[object, ...]]],
        proposed: list[tuple[str, tuple[object, ...]]],
    ) -> None:
        max_length = max(len(original), len(proposed))
        for index in range(max_length):
            location = original[index][0] if index < len(original) else proposed[index][0]
            before = original[index][1] if index < len(original) else ("missing",)
            after = proposed[index][1] if index < len(proposed) else ("missing",)
            if before != after:
                differences.append(
                    self._difference(f"{prefix}.{location}", location, before, after)
                )

    @staticmethod
    def _difference(
        location: str, field: str, original: object, proposed: object
    ) -> DetectedFormatDifference:
        category = (
            FormatDifferenceCategory.MARGIN
            if "margin" in field
            else FormatDifferenceCategory.LINE_SPACING
            if "line_spacing" in field
            else FormatDifferenceCategory.FONT_SIZE
            if "font.size" in field
            else FormatDifferenceCategory.COLOR
            if "color" in field
            else FormatDifferenceCategory.FONT
            if "font" in field
            else FormatDifferenceCategory.TABLE
            if "table" in location or "cells" in location
            else FormatDifferenceCategory.OTHER
        )
        return DetectedFormatDifference(
            category=category,
            location=location,
            original_value=str(original),
            proposed_value=str(proposed),
        )

    @staticmethod
    def _format_for_input_kind(input_kind: str | None) -> OriginalFormat | None:
        if input_kind == "editable_docx":
            return OriginalFormat.DOCX
        if input_kind == "text_pdf":
            return OriginalFormat.PDF
        return None

    @staticmethod
    def _value(value: object) -> object:
        if value is None:
            return None
        if isinstance(value, int | float):
            return float(value)
        return str(value)

    @staticmethod
    def _color_value(value: object) -> object:
        return str(value) if value is not None else None

    @staticmethod
    def _error_value(error: Exception) -> str:
        return f"{type(error).__name__}: {error}"[:512]

    @staticmethod
    def _has_no_text_spans(values: list[tuple[str, tuple[object, ...]]]) -> bool:
        return ("text_spans", ("missing",)) in values
