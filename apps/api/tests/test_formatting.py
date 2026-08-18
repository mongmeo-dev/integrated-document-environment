from io import BytesIO
from uuid import uuid4

import fitz
import pytest
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

from ide_api.domains.documents.models import DocumentVersion
from ide_api.domains.formatting.comparison import StorageFormatComparisonRunner
from ide_api.domains.formatting.models import ExternalEditResult, FormatCheck
from ide_api.domains.formatting.schemas import (
    ExternalEditResultStatus,
    FormatDifferenceCategory,
    OriginalFormat,
    VisualReviewStatus,
)
from ide_api.domains.formatting.service import FormattingService


class MemoryStorage:
    def __init__(self, objects: dict[str, bytes]) -> None:
        self.objects = objects

    def download(self, object_key: str) -> BytesIO:
        return BytesIO(self.objects[object_key])


def _docx_bytes(
    *, margin_inches: float = 1, size: int = 11, color: RGBColor | None = None, table: bool = False
) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(margin_inches)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = paragraph.add_run("formatted")
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = color or RGBColor(0, 0, 0)
    if table:
        document.add_table(rows=1, cols=2).cell(0, 0).text = "table"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _source(*, input_kind: str = "editable_docx") -> DocumentVersion:
    return DocumentVersion(
        id=uuid4(),
        document_id=uuid4(),
        original_filename="source.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size_bytes=1,
        sha256="0" * 64,
        object_key="source",
        created_by_id=uuid4(),
        input_kind=input_kind,
    )


def _result(*, object_key: str = "result", original_format: str = "docx") -> ExternalEditResult:
    return ExternalEditResult(
        id=uuid4(),
        document_id=uuid4(),
        document_version_id=uuid4(),
        original_format=original_format,
        original_filename="result.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size_bytes=1,
        sha256="1" * 64,
        object_key=object_key,
        created_by_id=uuid4(),
        status=ExternalEditResultStatus.UPLOADED.value,
    )


@pytest.mark.asyncio
async def test_docx_comparison_detects_font_size_color_table_and_margin_changes() -> None:
    source = _docx_bytes()
    result = _docx_bytes(margin_inches=1.5, size=14, color=RGBColor(255, 0, 0), table=True)
    runner = StorageFormatComparisonRunner(MemoryStorage({"source": source, "result": result}))

    differences = await runner.compare(original=_source(), result=_result())

    categories = {difference.category for difference in differences}
    assert FormatDifferenceCategory.MARGIN in categories
    assert FormatDifferenceCategory.FONT_SIZE in categories
    assert FormatDifferenceCategory.COLOR in categories
    assert FormatDifferenceCategory.TABLE in categories


@pytest.mark.asyncio
async def test_docx_comparison_returns_no_difference_for_identical_document() -> None:
    source = _docx_bytes(table=True)
    runner = StorageFormatComparisonRunner(MemoryStorage({"source": source, "result": source}))

    assert await runner.compare(original=_source(), result=_result()) == []


@pytest.mark.asyncio
async def test_cross_format_and_scanned_pdf_are_blocked() -> None:
    cross_format = await StorageFormatComparisonRunner(MemoryStorage({})).compare(
        original=_source(input_kind="text_pdf"),
        result=_result(original_format=OriginalFormat.DOCX.value),
    )
    scanned = fitz.open()
    scanned.new_page()
    scanned_bytes = scanned.tobytes()
    scanned.close()
    runner = StorageFormatComparisonRunner(
        MemoryStorage({"source": scanned_bytes, "result": scanned_bytes})
    )
    scan_differences = await runner.compare(
        original=_source(input_kind="text_pdf"),
        result=_result(original_format=OriginalFormat.PDF.value),
    )

    assert cross_format
    assert scan_differences
    assert any(
        difference.category is FormatDifferenceCategory.OTHER for difference in scan_differences
    )


@pytest.mark.asyncio
async def test_approval_requires_automatic_and_visual_passes() -> None:
    service = FormattingService(object(), StorageFormatComparisonRunner(MemoryStorage({})))
    result = _result()
    check = FormatCheck(
        automatic_check_completed=False,
        visual_review=VisualReviewStatus.PASSED.value,
        unresolved_difference_count=0,
    )
    result.status = ExternalEditResultStatus.PASSED.value
    result.format_check = check

    async def get_result(external_edit_result_id):
        return result

    service._get_external_edit_result = get_result  # type: ignore[method-assign]
    assert not await service.is_approval_allowed(external_edit_result_id=uuid4())
    check.automatic_check_completed = True
    check.visual_review = VisualReviewStatus.PENDING.value
    assert not await service.is_approval_allowed(external_edit_result_id=uuid4())
    check.visual_review = VisualReviewStatus.PASSED.value
    assert await service.is_approval_allowed(external_edit_result_id=uuid4())
