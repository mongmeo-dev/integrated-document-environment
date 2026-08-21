import hashlib
from io import BytesIO
from types import SimpleNamespace
from uuid import uuid4

import fitz
import pytest
from docx import Document as DocxDocument
from sqlalchemy import select, text

from ide_api.config import Settings
from ide_api.core.database import async_session
from ide_api.core.security import hash_password
from ide_api.domains.auth.models import User
from ide_api.domains.documents.extraction import (
    extract_document_text,
    render_scanned_pdf_pages,
)
from ide_api.domains.documents.models import Document, DocumentVersion
from ide_api.domains.evidence.models import DocumentEvidenceLink, EvidenceItem
from ide_api.domains.impacts.analysis import RelationshipAnalysisService
from ide_api.domains.impacts.models import DocumentRelationship, RelationshipAnalysisRun
from ide_api.infrastructure.openai import (
    DocumentAnalysisInput,
    EvidenceAnalysisInput,
    OpenAIRelationshipAnalysis,
    OpenAIRelationshipAnalyzer,
    RelationshipAnalysisResult,
)


class _FakeResponses:
    def __init__(self, parsed: RelationshipAnalysisResult) -> None:
        self.parsed = parsed
        self.request: dict[str, object] | None = None

    async def parse(self, **kwargs: object) -> SimpleNamespace:
        self.request = kwargs
        return SimpleNamespace(output_parsed=self.parsed, model=kwargs["model"])


class _FakeOpenAI:
    def __init__(self, parsed: RelationshipAnalysisResult) -> None:
        self.responses = _FakeResponses(parsed)


class _FakeStorage:
    def __init__(self, objects: dict[str, bytes]) -> None:
        self.objects = objects

    def download(self, object_key: str) -> BytesIO:
        return BytesIO(self.objects[object_key])


class _FakeAnalyzer:
    def __init__(self, result: RelationshipAnalysisResult) -> None:
        self.result = result

    def route_model(self, **_: object) -> str:
        return "routed-model"

    async def analyze(self, **_: object) -> OpenAIRelationshipAnalysis:
        return OpenAIRelationshipAnalysis(
            result=self.result,
            model_id="routed-model",
            prompt_version="relationship-analysis-v1",
        )


def _settings() -> Settings:
    return Settings(
        _env_file=None,
        OPENAI_API_KEY="test-key",
        openai_relationship_model_standard="standard-model",
        openai_relationship_model_complex="complex-model",
    )


@pytest.mark.asyncio
async def test_analyzer_routes_small_structured_analysis_to_standard_model() -> None:
    target_id = uuid4()
    evidence_id = uuid4()
    parsed = RelationshipAnalysisResult.model_validate(
        {
            "document_relationships": [
                {
                    "target_document_id": str(target_id),
                    "source_location": "section 2",
                    "target_location": "requirements table",
                    "relationship_type": "semantic",
                    "reason": "Both specify the same retention requirement.",
                }
            ],
            "evidence_relationships": [
                {
                    "evidence_id": str(evidence_id),
                    "reason": "The cloud setting implements the documented retention period.",
                }
            ],
        }
    )
    client = _FakeOpenAI(parsed)
    analyzer = OpenAIRelationshipAnalyzer(settings=_settings(), client=client)  # type: ignore[arg-type]

    result = await analyzer.analyze(
        source=DocumentAnalysisInput(id=uuid4(), filename="new.tex", content="Retention: 7 years"),
        documents=[
            DocumentAnalysisInput(id=target_id, filename="policy.tex", content="Keep for 7 years")
        ],
        evidence=[
            EvidenceAnalysisInput(
                id=evidence_id,
                evidence_type="cloud_config",
                title="Object lifecycle",
                description="Retention is seven years.",
                reference="ncloud://bucket/policy",
                location="lifecycle rule",
                version="3",
            )
        ],
    )

    assert result.model_id == "standard-model"
    assert result.result.document_relationships[0].target_document_id == target_id
    assert client.responses.request is not None
    assert client.responses.request["model"] == "standard-model"
    assert client.responses.request["text_format"] is RelationshipAnalysisResult


@pytest.mark.asyncio
async def test_analyzer_routes_large_analysis_to_complex_model() -> None:
    client = _FakeOpenAI(
        RelationshipAnalysisResult(document_relationships=[], evidence_relationships=[])
    )
    analyzer = OpenAIRelationshipAnalyzer(settings=_settings(), client=client)  # type: ignore[arg-type]

    result = await analyzer.analyze(
        source=DocumentAnalysisInput(id=uuid4(), filename="large.tex", content="x" * 120_000),
        documents=[],
        evidence=[],
    )

    assert result.model_id == "complex-model"


def test_extracts_docx_paragraphs_and_tables() -> None:
    document = DocxDocument()
    document.add_paragraph("Server API requirement")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Endpoint"
    table.cell(0, 1).text = "/api/v1/health"
    content = BytesIO()
    document.save(content)

    extracted = extract_document_text(content.getvalue(), "requirements.docx")

    assert "Server API requirement" in extracted
    assert "Endpoint | /api/v1/health" in extracted


def test_scanned_pdf_is_rendered_for_vision_analysis() -> None:
    document = fitz.open()
    page = document.new_page()
    page.draw_rect(fitz.Rect(20, 20, 200, 100), color=(0, 0, 0), fill=(0.5, 0.5, 0.5))
    content = document.tobytes()
    document.close()

    assert "page images" in extract_document_text(content, "scan.pdf")
    images = render_scanned_pdf_pages(content)
    assert len(images) == 1
    assert images[0].startswith(b"\xff\xd8")


def test_relationship_analysis_batches_large_candidate_sets_without_dropping_items() -> None:
    documents = [
        DocumentAnalysisInput(id=uuid4(), filename=f"{index}.tex", content="x" * 100)
        for index in range(25)
    ]
    evidence = [
        EvidenceAnalysisInput(
            id=uuid4(),
            evidence_type="server_code",
            title=f"Service {index}",
            description="API implementation",
            reference=None,
            location=None,
            version=None,
        )
        for index in range(20)
    ]

    batches = RelationshipAnalysisService._analysis_batches(documents, evidence)

    assert all(
        len(document_batch) + len(evidence_batch) <= 20
        for document_batch, evidence_batch in batches
    )
    assert sum(len(document_batch) for document_batch, _ in batches) == len(documents)
    assert sum(len(evidence_batch) for _, evidence_batch in batches) == len(evidence)


@pytest.mark.asyncio
async def test_analysis_persists_unconfirmed_document_and_product_evidence_candidates() -> None:
    source_content = (
        b"\\documentclass{article}\\begin{document}"
        b"The API stores files for seven years.\\end{document}"
    )
    target_content = (
        b"\\documentclass{article}\\begin{document}"
        b"Retention policy: seven years.\\end{document}"
    )
    source = Document()
    target = Document()
    evidence = EvidenceItem(
        evidence_type="cloud_config",
        title="Object storage lifecycle",
        description="The production bucket retains objects for seven years.",
        reference="ncloud://production-documents",
        location="lifecycle rule",
        version="7",
    )

    async with async_session() as session:
        await session.execute(text("TRUNCATE TABLE users RESTART IDENTITY CASCADE"))
        user = User(
            email="analysis@neudive.com",
            display_name="분석 테스트",
            password_hash=hash_password("correct-horse-battery-staple"),
        )
        session.add(user)
        await session.flush()
        source.versions.append(
            DocumentVersion(
                original_filename="source.tex",
                media_type="text/x-tex",
                size_bytes=len(source_content),
                sha256=hashlib.sha256(source_content).hexdigest(),
                object_key="source",
                created_by_id=user.id,
                status="ready",
                input_kind="latex_project",
            )
        )
        target.versions.append(
            DocumentVersion(
                original_filename="target.tex",
                media_type="text/x-tex",
                size_bytes=len(target_content),
                sha256=hashlib.sha256(target_content).hexdigest(),
                object_key="target",
                created_by_id=user.id,
                status="ready",
                input_kind="latex_project",
            )
        )
        session.add_all([source, target, evidence])
        await session.commit()

        analyzer = _FakeAnalyzer(
            RelationshipAnalysisResult.model_validate(
                {
                    "document_relationships": [
                        {
                            "target_document_id": str(target.id),
                            "source_location": "retention requirement",
                            "target_location": "retention policy",
                            "relationship_type": "semantic",
                            "reason": "Both define a seven-year retention period.",
                        }
                    ],
                    "evidence_relationships": [
                        {
                            "evidence_id": str(evidence.id),
                            "reason": "The cloud lifecycle rule implements the requirement.",
                        }
                    ],
                }
            )
        )
        service = RelationshipAnalysisService(
            session,
            storage=_FakeStorage({"source": source_content, "target": target_content}),  # type: ignore[arg-type]
            analyzer=analyzer,  # type: ignore[arg-type]
            settings=_settings(),
        )

        queued_run_id = await service.queue_registered_document(document_id=source.id)
        completed_run_id = await service.analyze_registered_document(document_id=source.id)

        run = await session.get(RelationshipAnalysisRun, completed_run_id)
        relationship = (
            await session.execute(
                select(DocumentRelationship).where(
                    DocumentRelationship.analysis_run_id == completed_run_id
                )
            )
        ).scalar_one()
        evidence_link = (
            await session.execute(
                select(DocumentEvidenceLink).where(
                    DocumentEvidenceLink.analysis_run_id == completed_run_id
                )
            )
        ).scalar_one()

        assert completed_run_id == queued_run_id
        assert run is not None
        assert run.status == "completed"
        assert run.model_id == "routed-model"
        assert relationship.status == "candidate"
        assert relationship.target_document_id == target.id
        assert evidence_link.status == "candidate"
        assert evidence_link.evidence_id == evidence.id
